import os
import logging
from typing import List
from langchain_core.documents import Document
import pypdf
import docx

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def validate_file(file_path: str) -> None:
    """
    Validates that the file exists, is not empty, and has a supported extension.
    Raises FileNotFoundError or ValueError if validation fails.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found on system paths: {file_path}")
    
    file_size = os.path.getsize(file_path)
    if file_size == 0:
        raise ValueError(
            f"Validation Failed: The file '{os.path.basename(file_path)}' is empty (0 bytes)."
        )
    
    # Check extension
    _, ext = os.path.splitext(file_path.lower())
    supported_extensions = {'.pdf', '.docx', '.txt'}
    if ext not in supported_extensions:
        raise ValueError(
            f"Unsupported file format '{ext}'. Supported formats are: PDF, DOCX, TXT."
        )

def load_pdf(file_path: str) -> List[Document]:
    """
    Extracts text page-by-page from a PDF file using pypdf.
    Stores the filename and page numbers (1-indexed) in the metadata.
    """
    documents = []
    file_name = os.path.basename(file_path)
    
    try:
        reader = pypdf.PdfReader(file_path)
        
        # Check if PDF is encrypted/password-protected
        if reader.is_encrypted:
            raise ValueError(
                f"The PDF file '{file_name}' is encrypted or password-protected. "
                "Please decrypt the file and try again."
            )
            
        num_pages = len(reader.pages)
        if num_pages == 0:
            raise ValueError(f"The PDF file '{file_name}' contains no pages.")
            
        for page_num in range(num_pages):
            try:
                page = reader.pages[page_num]
                text = page.extract_text()
                
                # Only append if page has extractable text
                if text and text.strip():
                    documents.append(
                        Document(
                            page_content=text.strip(),
                            metadata={
                                "source": file_name,
                                "page": page_num + 1
                            }
                        )
                    )
            except Exception as page_err:
                logger.warning(
                    f"Error extracting page {page_num + 1} from {file_name}: {page_err}"
                )
                
        if not documents:
            raise ValueError(
                f"Could not extract any readable text from '{file_name}'. "
                "The file might consist entirely of scanned images/drawings. "
                "Ensure it contains selectable digital text."
            )
            
    except Exception as e:
        logger.error(f"Failed to read PDF file {file_name}: {e}")
        err_msg = str(e)
        if "crypt" in err_msg.lower() or "password" in err_msg.lower():
            raise RuntimeError(
                f"The PDF file '{file_name}' is password-protected and cannot be parsed."
            )
        else:
            raise RuntimeError(
                f"Failed to parse PDF file '{file_name}'. The file may be corrupt or formatted incorrectly. "
                f"Details: {err_msg}"
            )
        
    return documents

def load_docx(file_path: str) -> List[Document]:
    """
    Extracts text from a DOCX file paragraph by paragraph using python-docx.
    Stores the filename in the metadata.
    """
    documents = []
    file_name = os.path.basename(file_path)
    
    try:
        doc = docx.Document(file_path)
        full_text = []
        
        # Extract paragraph text
        for para in doc.paragraphs:
            if para.text and para.text.strip():
                full_text.append(para.text.strip())
                
        # Extract table text if available
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    full_text.append(" | ".join(row_text))
                    
        combined_text = "\n\n".join(full_text)
        
        if not combined_text.strip():
            raise ValueError(f"The DOCX file '{file_name}' contains no readable text content.")
            
        documents.append(
            Document(
                page_content=combined_text.strip(),
                metadata={
                    "source": file_name,
                    "page": "N/A"
                }
            )
        )
        
    except Exception as e:
        logger.error(f"Failed to read DOCX file {file_name}: {e}")
        raise RuntimeError(
            f"Failed to parse DOCX file '{file_name}'. The file may be corrupt or not a valid Word document. "
            f"Details: {str(e)}"
        )
        
    return documents

def load_txt(file_path: str) -> List[Document]:
    """
    Reads a plain text file using fallback encodings to handle special characters.
    Stores the filename in the metadata.
    """
    file_name = os.path.basename(file_path)
    encodings = ['utf-8', 'latin-1', 'cp1252']
    text_content = None
    
    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                text_content = f.read()
            break
        except UnicodeDecodeError:
            continue
            
    if text_content is None:
        raise RuntimeError(
            f"Failed to decode text file '{file_name}' using standard encodings (UTF-8, Latin-1, CP1252)."
        )
        
    if not text_content.strip():
        raise ValueError(f"The text file '{file_name}' is empty or contains only whitespace.")
        
    return [
        Document(
            page_content=text_content.strip(),
            metadata={
                "source": file_name,
                "page": "N/A"
            }
        )
    ]

def load_document(file_path: str) -> List[Document]:
    """
    Validates a file, detects its type, and loads it into a list of LangChain Document objects.
    Raises ValueError or RuntimeError on error.
    """
    validate_file(file_path)
    _, ext = os.path.splitext(file_path.lower())
    
    if ext == '.pdf':
        return load_pdf(file_path)
    elif ext == '.docx':
        return load_docx(file_path)
    elif ext == '.txt':
        return load_txt(file_path)
    else:
        raise ValueError(f"Unsupported file format: {ext}")
